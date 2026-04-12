# Функция для прочтения текст из файла a.docx (D:\KOVKA\2025\1.docx) и изменение z.docx с выводом результата в терминал. [26-04-12 - 15:03:56]

from datetime import datetime as dt

MTIME = dt.now().strftime('%Y-%m-%d - %H:%M:%S')

from docx import Document
def read_docx(fpath):
    doc = Document(fpath)
    сtext = []
    for paragraph in doc.paragraphs:
        сtext.append(paragraph.text)
    return '\n'.join(сtext)

# Функция которая заменяет все [n] на строку с индексом n в списке (список передается в эту функцию) [26-04-13 - 0:45:56]

from termcolor import colored
def replace_placeholders(сtext, replacements, placeholder_format='[{}]', color = None, log = True):
    for i, replacement in enumerate(replacements):
        placeholder = placeholder_format.format(i)
        сtext = сtext.replace(placeholder, replacement)
    if log and color:
        print(colored(f'\nЗамены выполнены успешно. [{MTIME}]', 'red'))

        return colored(сtext, color), сtext
    elif log:
        print(colored(f'\nЗамены выполнены успешно. [{MTIME}]', 'red'))
    elif color:
        return colored(сtext, color)
    else: return сtext

# Функция которая изменяет файл (путь к файлу передается в эту функцию) основываясь на result из 25 строки (результат замены) [26-04-12 - 15:12:56]

def write_docx(fpath, сtext, log = 1):
    doc = Document()
    for line in сtext.split('\n'):
        doc.add_paragraph(line)
    doc.save(fpath)
    if log:
        print(colored(f'\nФайл сохранен по пути: {fpath} [{MTIME}]', 'red'))

# Пример использования функций [26-04-13 - 0:45:56]

mlp, mpath      = [' m - main, l - list, p - parameter ', 'Ольги', 'Долгопрудный', '88005553535',
'Легенда', 'которая учится в лучшем ВУЗе Россиии - МФТИ'], '/home/codespace/etc/tmp/github/1/'
spath, epath    = mpath + 'a.docx', mpath + 'z.docx'
source, sep     = read_docx(spath), colored('-' * 67, 'red')
result          = replace_placeholders(source, mlp, color = 'green')
print(colored(source, 'magenta'), result[0], sep = f'\n\n{sep}\n\n')
write_docx(epath, result[1])

''' 
Ошибки при выполнении кода и их решения [26-04-12 - 15:03:56] 
'''

''' 
Traceback (most recent call last):
  File '/home/codespace/etc/tmp/github/1/main.py', line 2, in <module>
    from docx import Document
ModuleNotFoundError: No module named 'docx' 
'''

# Для решения этой проблемы необходимо установить библиотеку python-docx, которая позволяет работать с файлами .docx. Вы можете установить ее с помощью pip (pip install python-docx)
